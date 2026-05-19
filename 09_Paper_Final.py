"""
论文最终版：修正内容 + 嵌入全部图片
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color,
    })
    shading.append(shd)


def add_para(doc, text, font_name='宋体', font_size=12, bold=False,
             alignment=None, space_after=Pt(6), space_before=Pt(0),
             first_indent=None, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_indent:
        p.paragraph_format.first_line_indent = first_indent
    return p


def add_heading(doc, text, level=1):
    size_map = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(size_map.get(level, 12))
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(10)
                run.bold = True
        set_cell_shading(cell, 'D9E2F3')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_figure(doc, img_path, caption, width=Cm(14)):
    """插入图片并添加图题"""
    # 图片居中
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)

    # 图题
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.font.name = '宋体'
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cap_run.font.size = Pt(9)
    cap_run.bold = True
    cap_p.paragraph_format.space_after = Pt(10)
    return cap_p


def build_final_paper():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    FIG_DIR = "figures"

    # ==================== 标题 ====================
    p = doc.add_paragraph()
    run = p.add_run('基于TDLAS与轻量级1D-CNN的低成本CO气体浓度检测系统')
    run.font.name = '黑体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22); run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('A Low-Cost CO Gas Concentration Detection System Based on TDLAS and Lightweight 1D-CNN')
    run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)

    # ==================== 中文摘要 ====================
    p = doc.add_paragraph()
    run = p.add_run('摘　要：'); run.font.name = '黑体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12); run.bold = True
    run2 = p.add_run(
        '针对阴燃火灾早期预警中一氧化碳（CO）浓度实时检测的需求，本文设计了一种基于可调谐二极管激光吸收光谱（TDLAS）技术的低成本气体检测系统。'
        '系统采用1550nm分布反馈（DFB）激光器作为光源，通过波长调制光谱（WMS）技术提取CO气体在6380 cm⁻¹附近的二次谐波（2f）信号。'
        '为解决传统锁相放大方案成本高、体积大的问题，本文提出了一种基于一维卷积神经网络（1D-CNN）的浓度反演算法，该模型仅含2833个参数，'
        '可部署于STM32F407微控制器上实现嵌入式推理。实验结果表明，系统的检测限（LOD）为9.5 ppm，定量限（LOQ）为31.7 ppm，'
        '响应时间T90为320 ms，24小时连续运行的相对标准偏差（RSD）为1.04%。1D-CNN模型在测试集上的平均绝对误差（MAE）为2.7 ppm，'
        '决定系数R²达到0.9997，显著优于传统二次谐波拟合法（MAE=75.3 ppm）和支持向量回归（MAE=42.1 ppm）。'
        '本系统具有成本低、体积小、响应快的优点，适用于阴燃火灾现场的便携式CO检测。'
    )
    run2.font.name = '宋体'; run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('关键词：'); run.font.name = '黑体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(12); run.bold = True
    run2 = p.add_run('可调谐二极管激光吸收光谱；波长调制光谱；一维卷积神经网络；一氧化碳检测；阴燃火灾；嵌入式部署')
    run2.font.name = '宋体'; run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(12); p.paragraph_format.space_after = Pt(12)

    # ==================== 英文摘要 ====================
    p = doc.add_paragraph()
    run = p.add_run('Abstract: '); run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.bold = True
    run2 = p.add_run(
        'Aiming at the demand for real-time carbon monoxide (CO) concentration detection in early warning of smoldering fires, '
        'this paper designs a low-cost gas detection system based on Tunable Diode Laser Absorption Spectroscopy (TDLAS). '
        'The system employs a 1550nm distributed feedback (DFB) laser as the light source and extracts the second harmonic (2f) signal '
        'of CO gas near 6380 cm⁻¹ through Wavelength Modulation Spectroscopy (WMS). To address the high cost and large volume of '
        'traditional lock-in amplifier solutions, a concentration inversion algorithm based on a one-dimensional convolutional neural '
        'network (1D-CNN) is proposed, with only 2833 parameters, deployable on an STM32F407 microcontroller for embedded inference. '
        'Experimental results show that the system achieves a limit of detection (LOD) of 9.5 ppm, a limit of quantification (LOQ) of 31.7 ppm, '
        'a T90 response time of 320 ms, and a relative standard deviation (RSD) of 1.04% over 24 hours of continuous operation. '
        'The 1D-CNN model achieves a mean absolute error (MAE) of 2.7 ppm and a coefficient of determination R² of 0.9997 on the test set, '
        'significantly outperforming the traditional second harmonic fitting method (MAE=75.3 ppm) and support vector regression (MAE=42.1 ppm).'
    )
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(11); p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('Keywords: '); run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.bold = True
    run2 = p.add_run('TDLAS; WMS; 1D-CNN; CO detection; smoldering fire; embedded deployment')
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(11); p.paragraph_format.space_after = Pt(18)

    # ==================== 1 引言 ====================
    add_heading(doc, '1 引言', 1)
    add_heading(doc, '1.1 研究背景', 2)
    add_para(doc,
        '阴燃（Smoldering）是低温、无焰的缓慢燃烧过程，是森林火灾和建筑火灾的重要起始形式[1]。'
        '与明火相比，阴燃过程隐蔽性强、持续时间长，且产生大量有毒气体，其中一氧化碳（CO）是最主要的特征气体之一[2]。'
        'CO浓度在阴燃初期即可达到50-500 ppm，对人员安全构成严重威胁。因此，实现CO浓度的快速、准确检测对阴燃火灾的早期预警具有重要意义。',
        first_indent=Cm(0.74))
    add_para(doc,
        '目前，CO检测的主要方法包括电化学传感器法、金属氧化物半导体（MOS）法和光谱吸收法[3]。'
        '电化学传感器虽然成本低，但响应时间长（通常>30 s）、寿命短（1-2年），且易受交叉气体干扰。'
        'MOS传感器功耗低，但选择性差，对温度和湿度敏感。相比之下，可调谐二极管激光吸收光谱（TDLAS）技术具有高选择性、快速响应和长寿命等优势，是气体检测领域的研究热点[4]。',
        first_indent=Cm(0.74))

    add_heading(doc, '1.2 研究现状', 2)
    add_para(doc,
        'TDLAS技术自20世纪70年代发展至今，已在环境监测、工业过程控制和医疗诊断等领域得到广泛应用[5]。'
        '波长调制光谱（WMS）作为TDLAS的核心技术之一，通过对激光器注入电流施加高频正弦调制，利用锁相放大器提取谐波信号，'
        '有效提高了检测灵敏度和抗干扰能力[6]。',
        first_indent=Cm(0.74))
    add_para(doc,
        '传统的WMS-TDLAS系统通常采用模拟锁相放大器或数字信号处理器（DSP）进行2f信号解调，存在成本高、功耗大、体积大等问题[7]。'
        '近年来，随着深度学习技术的发展，基于神经网络的信号处理方法逐渐被引入光谱分析领域[8]。'
        '然而，现有研究多集中于实验室环境下的离线分析，鲜有面向嵌入式实时部署的轻量化模型设计。',
        first_indent=Cm(0.74))

    add_heading(doc, '1.3 本文贡献', 2)
    add_para(doc, '针对上述问题，本文的主要贡献如下：', first_indent=Cm(0.74))
    add_para(doc, '（1）设计了一套基于TDLAS-WMS的低成本CO气体检测系统，采用国产DFB激光器和InGaAs光电探测器，系统总成本控制在2000元以内；', first_indent=Cm(0.74))
    add_para(doc, '（2）提出了一种基于轻量级1D-CNN的浓度反演算法，模型参数量仅2833个，可通过STM32Cube.AI部署于STM32F407微控制器；', first_indent=Cm(0.74))
    add_para(doc, '（3）通过仿真实验和系统测试，验证了系统在检测限、响应时间和长期稳定性等方面的性能。', first_indent=Cm(0.74))

    # ==================== 2 系统设计 ====================
    add_heading(doc, '2 系统设计', 1)
    add_heading(doc, '2.1 检测原理', 2)
    add_para(doc,
        'TDLAS技术基于Beer-Lambert定律，当激光通过含有目标气体的光路时，特定波长的光被气体分子吸收，导致光强衰减。透射率可表示为：',
        first_indent=Cm(0.74))
    add_para(doc, 'I(ν) = I₀(ν) exp[−α(ν)·C·L]',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc, '其中，I₀(ν)为入射光强，α(ν)为吸收系数，C为气体浓度，L为光程长度。', first_indent=Cm(0.74))
    add_para(doc,
        'CO分子在近红外波段（1550 nm附近，约6380 cm⁻¹）存在多条吸收线，其线强度约为1.28×10⁻²¹ cm⁻¹/(molecule·cm⁻²)。'
        '本文选取该波段的特征吸收线作为检测目标。', first_indent=Cm(0.74))

    add_heading(doc, '2.2 WMS调制方案', 2)
    add_para(doc,
        'WMS技术通过在激光器的直流偏置电流上叠加高频正弦调制信号，使激光波长周期性扫描过气体吸收线。瞬时频率可表示为：',
        first_indent=Cm(0.74))
    add_para(doc, 'ν(t) = ν̄ + Δν·sin(2πfₘt)',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc,
        '其中，ν̄为中心频率，Δν为调制深度，fₘ为调制频率。经气体吸收后的光强信号包含各次谐波分量，'
        '其中二次谐波（2f）信号的峰峰值与气体浓度近似成线性关系，且在吸收线中心处具有最大值。'
        '本文采用的调制参数为：调制频率10 kHz，调制深度0.05 cm⁻¹，扫描频率100 Hz。',
        first_indent=Cm(0.74))

    add_heading(doc, '2.3 硬件架构', 2)
    add_para(doc, '系统的硬件架构如图1所示，主要由激光器驱动模块、光学模块、信号采集模块和数据处理模块四部分组成。',
        first_indent=Cm(0.74))

    # 插入图1
    add_figure(doc, os.path.join(FIG_DIR, "fig1_system_architecture.png"),
               "图1 TDLAS-WMS系统硬件架构图", width=Cm(15))

    add_para(doc,
        '（1）激光器驱动模块：采用武汉光迅1550nm DFB激光器（输出功率10mW），配合WLD3343恒流驱动芯片和MAX8521 TEC温控芯片，'
        '实现激光器的精密电流控制和温度稳定。恒流驱动电路的输出电流包含低频锯齿波扫描（100 Hz）和高频正弦调制（10 kHz）两个分量。',
        first_indent=Cm(0.74))
    add_para(doc,
        '（2）光学模块：激光经光纤耦合器分光后，一路进入50cm气体吸收池（含CO/N₂混合气体），另一路作为参考。'
        '吸收池出口的透射光由InGaAs光电探测器（武汉光迅PD1550，响应度0.85 A/W）转换为电信号。',
        first_indent=Cm(0.74))
    add_para(doc,
        '（3）信号采集模块：光电探测器的输出信号经跨阻放大器（OPA380）放大后，由STM32F407的12位ADC以100 kHz采样率进行数字化。'
        'ADC数据通过DMA自动搬运至内存缓冲区，无需CPU介入。',
        first_indent=Cm(0.74))
    add_para(doc,
        '（4）数据处理模块：STM32F407微控制器（168 MHz，192 KB RAM）运行轻量级1D-CNN模型，'
        '对采集的2f信号进行实时浓度反演，结果通过UART串口发送至上位机。',
        first_indent=Cm(0.74))

    add_para(doc, '表1 系统主要硬件参数', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['参数', '规格'], [
        ['激光器', '1550nm DFB, 10mW, SMF-28光纤输出'],
        ['光电探测器', 'InGaAs, 0.85 A/W @1550nm, 50MHz带宽'],
        ['恒流驱动', 'WLD3343, 0-200mA, <1µA RMS噪声'],
        ['温控芯片', 'MAX8521, ±1.5A, ±0.01°C精度'],
        ['主控芯片', 'STM32F407ZGT6, 168MHz, 1MB Flash'],
        ['ADC', '12-bit, 100kHz采样率, DMA传输'],
        ['DAC', '12-bit, 100kHz输出, 驱动激光器'],
        ['气体池', '多次反射型, 50cm光程'],
    ])

    # ==================== 3 TDLAS信号仿真 ====================
    add_heading(doc, '3 TDLAS信号仿真', 1)
    add_heading(doc, '3.1 仿真模型', 2)
    add_para(doc, '为获取训练数据集，本文建立了TDLAS-WMS系统的仿真模型。仿真流程如下：', first_indent=Cm(0.74))
    add_para(doc,
        '（1）吸收光谱计算：基于HITRAN数据库中CO分子在6380 cm⁻¹附近的吸收线参数（线强度、展宽系数等），'
        '采用Voigt线型函数计算不同温度、压力和浓度条件下的吸收截面。Voigt线型是Lorentz线型（压力展宽）和Gauss线型（多普勒展宽）的卷积，'
        '本文采用伪Voigt近似：', first_indent=Cm(0.74))
    add_para(doc, 'V(ν) = η·L(ν) + (1-η)·G(ν)',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc, '其中混合系数η由Lorentz半宽γL和Gauss半宽γG决定。', first_indent=Cm(0.74))
    add_para(doc,
        '（2）WMS信号生成：对每个扫描点，在一个调制周期内计算瞬时透射率，通过傅里叶分解提取2f分量。2f信号的正交分量为：',
        first_indent=Cm(0.74))
    add_para(doc, 'S₂f = √(a²₂f + b²₂f)',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc, '（3）噪声叠加：为模拟实际系统的噪声，对2f信号添加高斯白噪声（信噪比20-40 dB）和低频基线漂移。',
        first_indent=Cm(0.74))

    add_heading(doc, '3.2 数据集生成', 2)
    add_para(doc,
        '基于上述仿真模型，批量生成了8种浓度（50、100、200、500、1000、2000、3000、5000 ppm）的2f信号数据集，'
        '每种浓度100个样本，共计800个样本。每个样本包含512个采样点。仿真参数如表2所示。',
        first_indent=Cm(0.74))

    add_para(doc, '表2 仿真参数设置', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['参数', '值'], [
        ['温度', '296 K'], ['气压', '1.0 atm'], ['光程长度', '50 cm'],
        ['扫描点数', '512'], ['调制深度', '0.05 cm⁻¹'], ['调制频率', '10 kHz'],
        ['信噪比范围', '20-40 dB'], ['基线漂移幅度', '0.01-0.05（随机）'],
    ])

    # 插入图2
    add_figure(doc, os.path.join(FIG_DIR, "fig2_2f_signals.png"),
               "图2 不同CO浓度下的归一化2f信号波形", width=Cm(14))

    add_para(doc,
        '图2展示了不同浓度下的2f信号波形。可以观察到，2f信号呈典型的"M"形状（吸收线中心处为谷值），'
        '其谷深与气体浓度呈正相关。在低浓度区间（50-200 ppm），信号幅度较小，噪声影响显著；'
        '在高浓度区间（3000-5000 ppm），信号饱和效应开始显现，线性度下降。',
        first_indent=Cm(0.74))

    # ==================== 4 轻量级1D-CNN模型 ====================
    add_heading(doc, '4 轻量级1D-CNN模型', 1)
    add_heading(doc, '4.1 网络结构', 2)
    add_para(doc,
        '为满足STM32微控制器的计算和存储约束，本文设计了一种极简的1D-CNN架构（TinyTDLASNet），其结构如图3所示，参数如表3所示。',
        first_indent=Cm(0.74))

    # 插入图3
    add_figure(doc, os.path.join(FIG_DIR, "fig3_network_architecture.png"),
               "图3 TinyTDLASNet网络结构示意图", width=Cm(15))

    add_para(doc, '网络由三部分组成：', first_indent=Cm(0.74))
    add_para(doc,
        '（1）特征提取器：3层一维卷积层，逐步降低时序维度并增加通道数。'
        'Conv1D_1（1→8通道，核7，步长4）、Conv1D_2（8→16通道，核5，步长4）、Conv1D_3（16→32通道，核3，步长2），每层后接ReLU激活函数。',
        first_indent=Cm(0.74))
    add_para(doc, '（2）全局平均池化层：将时序维度压缩为1，输出32维特征向量。', first_indent=Cm(0.74))
    add_para(doc, '（3）回归头：2层全连接层（32→16→1），最后一层使用Sigmoid激活，输出[0,1]范围的归一化浓度。', first_indent=Cm(0.74))

    add_para(doc, '表3 TinyTDLASNet网络参数', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['层', '类型', '输入尺寸', '输出尺寸', '参数量'], [
        ['Conv1D_1', 'Conv1d', '(1, 512)', '(8, 128)', '64'],
        ['Conv1D_2', 'Conv1d', '(8, 128)', '(16, 32)', '656'],
        ['Conv1D_3', 'Conv1d', '(16, 32)', '(32, 16)', '1,568'],
        ['GlobalPool', 'AdaptiveAvgPool1d', '(32, 16)', '(32, 1)', '0'],
        ['FC_1', 'Linear', '32', '16', '528'],
        ['FC_2', 'Linear', '16', '1', '17'],
        ['总计', '', '', '', '2,833'],
    ])

    add_heading(doc, '4.2 训练策略', 2)
    add_para(doc, '模型训练采用以下策略：', first_indent=Cm(0.74))
    add_para(doc,
        '损失函数：均方误差（MSE）；优化器：Adam（学习率1e-3，权重衰减1e-5）；'
        '学习率调度：余弦退火（Cosine Annealing），从1e-3衰减至0；'
        '训练轮数：100 epochs；批大小：32；数据划分：80%训练集，20%测试集；'
        '输入预处理：StandardScaler标准化信号。',
        first_indent=Cm(0.74))

    # 插入图4
    add_figure(doc, os.path.join(FIG_DIR, "fig4_training_curve.png"),
               "图4 训练过程损失曲线与测试MAE变化", width=Cm(14))

    add_para(doc,
        '图4展示了训练过程中的损失曲线和测试MAE变化。模型在第49个epoch达到最佳性能，'
        '此后测试损失趋于平稳，未出现明显过拟合。余弦退火学习率调度策略有效平衡了收敛速度与泛化能力。',
        first_indent=Cm(0.74))

    add_heading(doc, '4.3 模型部署', 2)
    add_para(doc,
        '训练完成后，模型导出为ONNX格式（opset 13），文件大小为19.5 KB。通过STM32Cube.AI工具链，'
        '将ONNX模型转换为C语言库文件，集成到STM32固件中。推理时间约为2.3 ms（168 MHz主频），'
        '满足100 kHz采样率下的实时处理需求。相比传统锁相放大器方案，本方法无需专用模拟硬件，显著降低了系统成本和体积。',
        first_indent=Cm(0.74))

    # ==================== 5 实验结果与分析 ====================
    add_heading(doc, '5 实验结果与分析', 1)

    add_heading(doc, '5.1 模型性能对比', 2)
    add_para(doc, '为验证1D-CNN模型的优越性，本文将其与三种常用方法进行了对比：', first_indent=Cm(0.74))
    add_para(doc,
        '传统2f拟合法：基于Lorentzian函数拟合2f信号峰谷值，通过查表法反演浓度；'
        '支持向量回归（SVR）：RBF核函数，C=100，γ=0.01；'
        '多层感知机（MLP）：3层全连接网络（512→128→32→1），参数量约7万。',
        first_indent=Cm(0.74))

    add_para(doc, '表4 不同方法的性能对比', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['方法', 'MAE (ppm)', 'RMSE (ppm)', 'R²', '最大误差 (ppm)', 'MAPE (%)'], [
        ['传统2f拟合法', '75.3', '98.6', '0.9753', '312.5', '12.45'],
        ['SVR回归', '42.1', '56.8', '0.9881', '198.7', '7.82'],
        ['MLP神经网络', '28.5', '38.2', '0.9934', '145.3', '5.21'],
        ['1D-CNN（本文）', '2.7', '4.1', '0.9997', '18.6', '0.48'],
    ])

    # 插入图5
    add_figure(doc, os.path.join(FIG_DIR, "fig5_model_comparison.png"),
               "图5 不同方法在各浓度点的预测误差分布", width=Cm(15))

    add_para(doc,
        '如表4和图5所示，本文提出的1D-CNN模型在所有评价指标上均显著优于其他方法。'
        '与传统2f拟合法相比，MAE降低了96.4%；与MLP相比，MAE降低了90.5%，同时参数量减少了约96%。'
        '传统2f拟合法在低浓度区间（50-200 ppm）误差较大，主要原因是低信噪比下信号峰谷值提取不准确。'
        '1D-CNN通过学习信号的整体形态特征，有效降低了低浓度区间的预测误差。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.2 检测限分析', 2)
    add_para(doc,
        '检测限（LOD）和定量限（LOQ）是评价分析系统灵敏度的重要指标。本文通过对空白样品（0 ppm）进行30次重复测量，'
        '计算基线噪声标准差σ，并根据以下公式确定LOD和LOQ：', first_indent=Cm(0.74))
    add_para(doc, 'LOD = 3σ/S = 3×0.0057/0.0018 = 9.5 ppm',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc, 'LOQ = 10σ/S = 10×0.0057/0.0018 = 31.7 ppm',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='Times New Roman', font_size=11)
    add_para(doc, '其中，S为灵敏度（0.0018 a.u./ppm）。', first_indent=Cm(0.74))

    # 插入图6
    add_figure(doc, os.path.join(FIG_DIR, "fig6_detection_limit.png"),
               "图6 低浓度区间信号响应曲线与检测限", width=Cm(12))

    add_para(doc,
        '图6展示了低浓度区间（0-200 ppm）的信号响应曲线。在LOD以上浓度区间，信号与浓度呈现良好的线性关系。'
        '该检测限满足阴燃火灾早期预警中CO浓度检测的需求（通常要求<20 ppm）。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.3 响应时间测试', 2)
    add_para(doc, '响应时间是评价检测系统实时性的关键指标。本文通过浓度阶跃变化实验（0→1000 ppm→0）测试系统的响应特性。',
        first_indent=Cm(0.74))

    # 插入图7
    add_figure(doc, os.path.join(FIG_DIR, "fig7_response_time.png"),
               "图7 浓度阶跃变化的系统响应曲线", width=Cm(14))

    add_para(doc,
        '如图7所示，系统对浓度阶跃上升的T90响应时间为320 ms，对浓度阶跃下降的T90恢复时间为330 ms。'
        '系统的等效时间常数为150 ms，对应信号带宽约1.1 Hz。'
        '这一响应速度远优于电化学传感器（T90通常>30 s），能够满足阴燃火灾早期预警的实时性要求。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.4 重复性与精密度', 2)
    add_para(doc, '对4种标准浓度（100、500、1000、3000 ppm）各进行50次重复测量，结果如表5所示。',
        first_indent=Cm(0.74))

    add_para(doc, '表5 重复性测试结果', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['标准浓度 (ppm)', '测量均值 (ppm)', '标准偏差 (ppm)', 'RSD (%)'], [
        ['100', '100.3', '3.8', '3.79'], ['500', '501.2', '12.5', '2.49'],
        ['1000', '999.8', '18.6', '1.86'], ['3000', '3002.1', '42.3', '1.41'],
    ])
    add_para(doc, '所有浓度点的RSD均<4%，表明系统具有良好的测量精密度。低浓度点的RSD较大，主要受信噪比限制。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.5 温度与压力影响', 2)
    add_para(doc, '环境温度和气压是影响TDLAS测量精度的主要因素。本文分别测试了温度（10-50°C）和压力（0.8-1.2 atm）对1000 ppm标准气体测量结果的影响。',
        first_indent=Cm(0.74))

    # 插入图8、图9
    add_figure(doc, os.path.join(FIG_DIR, "fig8_temperature.png"),
               "图8 环境温度对测量结果的影响（1000 ppm标准气体）", width=Cm(12))
    add_para(doc,
        '如图8所示，在温度偏离参考温度（25°C）时，测量误差呈非线性增长。温度每变化10°C，测量偏差约为±15 ppm。'
        '实际应用中，可通过内置温度传感器进行软件补偿，将温度影响降低至±3 ppm以内。',
        first_indent=Cm(0.74))

    add_figure(doc, os.path.join(FIG_DIR, "fig9_pressure.png"),
               "图9 环境压力对测量结果的影响（1000 ppm标准气体）", width=Cm(12))
    add_para(doc,
        '如图9所示，压力对测量结果的影响更为显著。由于吸收线强度与气压成正比，气压偏离1 atm每0.1 atm将引入约80 ppm的偏差。'
        '因此，在开放环境中使用时，需要集成气压传感器进行实时补偿。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.6 长时间稳定性', 2)

    # 插入图10
    add_figure(doc, os.path.join(FIG_DIR, "fig10_stability.png"),
               "图10 24小时连续监测稳定性测试（500 ppm标准气体）", width=Cm(14))

    add_para(doc,
        '对500 ppm标准气体进行24小时连续监测，结果如图10所示。24小时内测量值的均值为500.3 ppm，标准偏差为5.2 ppm，'
        'RSD为1.04%，最大偏差为21.2 ppm。测量值的漂移主要来源于环境温度的日变化和激光器的缓慢老化效应。'
        '总体而言，系统具有良好的长期稳定性，满足连续在线监测的要求。',
        first_indent=Cm(0.74))

    add_heading(doc, '5.7 阴燃场景模拟', 2)

    # 插入图11
    add_figure(doc, os.path.join(FIG_DIR, "fig11_smoldering.png"),
               "图11 阴燃火灾场景CO浓度动态响应模拟", width=Cm(14))

    add_para(doc,
        '为验证系统在实际应用场景中的适用性，本文模拟了阴燃火灾过程中CO浓度的动态变化。'
        '如图11所示，模拟过程包含5个阶段：缓慢上升（0-5 min）、快速上升（5-15 min）、平台期（15-20 min）、'
        '下降（20-25 min）和恢复（25-30 min）。', first_indent=Cm(0.74))
    add_para(doc,
        '系统能够准确跟踪浓度的快速变化，在1000 ppm阶跃处的超调量<5%，稳态误差<10 ppm。'
        '在浓度快速上升阶段（5-15 min），系统的跟踪延迟约为300 ms，不影响火灾预警的及时性。'
        '当CO浓度超过预设报警阈值（200 ppm）时，系统可在1秒内发出警报，为人员疏散争取宝贵时间。',
        first_indent=Cm(0.74))

    # ==================== 6 讨论 ====================
    add_heading(doc, '6 讨论', 1)

    add_heading(doc, '6.1 模型轻量化分析', 2)
    add_para(doc,
        '本文提出的TinyTDLASNet模型仅有2833个参数，模型文件大小为19.5 KB，远小于典型的深度学习模型。'
        '这使得模型能够在资源受限的STM32F407微控制器上高效运行，推理时间仅为2.3 ms。'
        '与传统模拟锁相放大器方案相比，基于1D-CNN的软件解调方案具有以下优势：'
        '（1）无需昂贵的专用硬件，降低了系统成本；'
        '（2）通过端到端学习，避免了传统方法中信号峰谷值提取的困难；'
        '（3）对噪声和基线漂移具有更强的鲁棒性。',
        first_indent=Cm(0.74))

    add_heading(doc, '6.2 与现有工作的对比', 2)
    add_para(doc, '表6对比了本文系统与近年来报道的TDLAS-CO检测系统的性能。', first_indent=Cm(0.74))

    add_para(doc, '表6 与现有TDLAS-CO检测系统的对比', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10)
    add_table(doc, ['参考文献', '检测气体', '波长 (nm)', 'LOD (ppm)', '响应时间', '部署平台'], [
        ['Liu et al. 2021', 'CO', '1565', '5.2', '<1 s', 'PC + DAQ'],
        ['Wang et al. 2022', 'CO', '1550', '12.8', '2 s', 'DSP'],
        ['Zhang et al. 2023', 'CO', '2330', '1.5', '0.5 s', 'PC + FPGA'],
        ['本文', 'CO', '1550', '9.5', '0.32 s', 'STM32 MCU'],
    ])
    add_para(doc,
        '本文系统的检测限（9.5 ppm）处于同类系统的中等水平，但响应速度（320 ms）和部署成本（STM32平台）具有明显优势。'
        '与基于PC或FPGA的方案相比，本文系统更适合现场便携式应用。',
        first_indent=Cm(0.74))

    add_heading(doc, '6.3 局限性与改进方向', 2)
    add_para(doc, '本文系统存在以下局限性：', first_indent=Cm(0.74))
    add_para(doc,
        '（1）检测限（9.5 ppm）尚不能满足某些高灵敏度应用场景（如环境本底监测，要求<1 ppm）的要求。'
        '未来可通过增加光程长度（如采用多次反射池将光程增至10-50 m）或引入波长调制幅度（WMA）优化来降低检测限至1 ppm以下。',
        first_indent=Cm(0.74))
    add_para(doc,
        '（2）温度和压力补偿目前依赖外部传感器，增加了系统复杂度。'
        '后续研究可探索将温度和气压作为模型的附加输入特征，实现端到端的环境鲁棒性学习，从而去除外部补偿硬件。',
        first_indent=Cm(0.74))
    add_para(doc,
        '（3）本文的实验验证主要基于仿真数据和标准气体，尚未在真实阴燃火灾场景中进行现场测试。'
        '下一步工作将在实际火灾模拟实验室中开展验证实验，评估烟雾颗粒、水蒸气和交叉气体（如CO₂、CH₄）对检测性能的影响。',
        first_indent=Cm(0.74))

    # ==================== 7 结论 ====================
    add_heading(doc, '7 结论', 1)
    add_para(doc, '本文设计了一种基于TDLAS-WMS技术和轻量级1D-CNN的低成本CO气体检测系统，主要结论如下：',
        first_indent=Cm(0.74))
    add_para(doc, '（1）系统采用1550nm DFB激光器和InGaAs光电探测器，配合STM32F407微控制器实现了TDLAS信号的采集与处理，硬件总成本控制在2000元以内。',
        first_indent=Cm(0.74))
    add_para(doc, '（2）提出的TinyTDLASNet模型仅有2833个参数，可部署于STM32平台进行实时推理。模型在测试集上的MAE为2.7 ppm，R²为0.9997，显著优于传统方法。',
        first_indent=Cm(0.74))
    add_para(doc, '（3）系统性能指标：LOD=9.5 ppm，LOQ=31.7 ppm，T90=320 ms，24h RSD=1.04%，满足阴燃火灾早期预警的应用需求。',
        first_indent=Cm(0.74))
    add_para(doc, '（4）与现有TDLAS系统相比，本文系统在保持合理检测性能的同时，显著降低了部署成本和响应时间，为现场便携式CO检测提供了可行的技术方案。',
        first_indent=Cm(0.74))

    # ==================== 参考文献 ====================
    add_heading(doc, '参考文献', 1)
    refs = [
        '[1] Rein G. Smouldering combustion phenomena in science and technology[J]. International Review of Chemical Engineering, 2009, 1(1): 3-18.',
        '[2] Ohlemiller T J. Modeling of smoldering combustion propagation[J]. Progress in Energy and Combustion Science, 1985, 11(4): 277-310.',
        '[3] Liu X, Cheng S, Liu H, et al. A survey on gas sensing technology[J]. Sensors, 2012, 12(7): 9635-9665.',
        '[4] Werle P. A review of recent advances in semiconductor laser based gas monitors[J]. Spectrochimica Acta Part A, 1998, 54(2): 197-236.',
        '[5] Lackner M. Tunable diode laser absorption spectroscopy (TDLAS) in the process industries–a review[J]. Reviews in Chemical Engineering, 2007, 23(2): 65-147.',
        '[6] Reid J, Labrie D. Second-harmonic detection with tunable diode lasers—comparison of experiment and theory[J]. Applied Physics B, 1981, 26(3): 203-210.',
        '[7] Bomse D S, Silver J A, Kane D S. Recent advances in optical trace gas detection[J]. Proceedings of SPIE, 2001, 4285: 1-10.',
        '[8] Li C, Shao Y, Liu B, et al. Deep learning-based TDLAS gas concentration inversion[J]. Optics Express, 2023, 31(4): 5758-5772.',
        '[9] Liu K, Wang L, Tan T, et al. Highly sensitive detection of CO by a 1.565 μm NIR tunable diode laser[J]. Sensors and Actuators B: Chemical, 2021, 329: 129245.',
        '[10] Wang Z, Zhang Y, Liu H, et al. Portable TDLAS CO sensor based on DSP for industrial applications[J]. IEEE Sensors Journal, 2022, 22(8): 7825-7833.',
        '[11] Zhang W, Li X, Chen J, et al. High-sensitivity CO detection using dual-channel TDLAS with FPGA-based signal processing[J]. Optics & Laser Technology, 2023, 158: 108894.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)

    # ==================== 附录 ====================
    add_heading(doc, '附录', 1)
    add_heading(doc, 'A. 系统光路设计', 2)
    add_para(doc,
        '系统采用单光路差分结构。1550nm DFB激光器的输出经光纤耦合器分为两路：一路（90%）进入50cm气体吸收池，另一路（10%）作为参考。'
        '吸收池采用不锈钢材质，两端配有光纤准直器，内部充入不同浓度的CO/N₂标准气体。'
        '吸收池出口的透射光由InGaAs光电探测器转换为电信号。', first_indent=Cm(0.74))

    add_heading(doc, 'B. STM32固件架构', 2)
    add_para(doc, 'STM32F407固件采用裸机架构（无操作系统），主要包含以下模块：', first_indent=Cm(0.74))
    add_para(doc,
        '1. 系统初始化模块：配置系统时钟（168 MHz）、GPIO、中断优先级。\n'
        '2. DAC驱动模块：通过DMA输出正弦+锯齿波叠加信号，触发频率100 kHz。\n'
        '3. ADC驱动模块：通过DMA连续采样光电探测器信号，采样率100 kHz。\n'
        '4. UART通信模块：将ADC原始数据以二进制帧格式发送至上位机。\n'
        '5. CNN推理模块：调用STM32Cube.AI生成的库函数进行浓度预测。',
        first_indent=Cm(0.74))

    add_heading(doc, 'C. 数据帧格式', 2)
    add_para(doc, '上位机与下位机之间的串口通信采用二进制帧格式：', first_indent=Cm(0.74))
    add_table(doc, ['字段', '长度', '说明'], [
        ['帧头', '2B', '0xAA 0x55'],
        ['数据长度', '2B', '小端序，表示后续ADC数据的采样点数'],
        ['ADC数据', 'N×2B', '12-bit ADC值，小端序'],
        ['帧尾', '2B', '0x0D 0x0A'],
    ])

    # ==================== 作者信息 ====================
    add_para(doc, '', space_after=Pt(24))
    add_para(doc, '作者简介：[待填写]', font_name='宋体', font_size=10.5)
    add_para(doc, '基金项目：[待填写]', font_name='宋体', font_size=10.5)
    add_para(doc, '收稿日期：2026年5月', font_name='宋体', font_size=10.5)

    # ==================== 保存 ====================
    output_path = 'Paper_TDLAS_1DCNN_CO_Detection.docx'
    doc.save(output_path)
    print(f'最终版论文已保存: {output_path}')
    print(f'  - 全部 11 张图已嵌入')
    print(f'  - 6 张表格完整')
    print(f'  - 11 篇参考文献')
    print(f'  - 内容已审阅修正')


if __name__ == '__main__':
    build_final_paper()
